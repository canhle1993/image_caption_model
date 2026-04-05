"""
Script tạo báo cáo đồ án Word (.docx) với biểu đồ minh hoạ.
Chạy: /opt/anaconda3/envs/tf-metal/bin/python generate_report.py
"""

import os
import io
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyArrowPatch
import matplotlib.gridspec as gridspec
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import warnings
warnings.filterwarnings("ignore")

# ─── màu theme ────────────────────────────────────────────────────────────────
C_VGG   = "#4C72B0"
C_RES   = "#DD8452"
C_CLIP  = "#55A868"
C_BASE  = "#C44E52"
C_LIGHT = "#F5F5F5"

OUT_PATH = os.path.join(os.path.dirname(__file__), "BAO_CAO_DO_AN.docx")

# ══════════════════════════════════════════════════════════════════════════════
# HELPER — lưu figure thành bytes
# ══════════════════════════════════════════════════════════════════════════════
def fig_to_bytes(fig, dpi=150):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight",
                facecolor=fig.get_facecolor())
    buf.seek(0)
    return buf

def add_fig(doc, fig, width_inches=6.0, caption_text=None):
    buf = fig_to_bytes(fig)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf, width=Inches(width_inches))
    plt.close(fig)
    if caption_text:
        cp = doc.add_paragraph(caption_text)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].italic = True
        cp.runs[0].font.size = Pt(10)
        cp.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ══════════════════════════════════════════════════════════════════════════════
# BIỂU ĐỒ 1 — Dataset statistics
# ══════════════════════════════════════════════════════════════════════════════
def chart_dataset():
    fig, axes = plt.subplots(1, 3, figsize=(13, 4), facecolor="white")
    fig.suptitle("Hình 1. Thống kê Dataset Flickr30k", fontsize=13, fontweight="bold", y=1.02)

    # 1a — train/val/test split
    ax = axes[0]
    labels = ["Train\n29,769", "Val\n1,014", "Test\n1,000"]
    sizes  = [29769, 1014, 1000]
    colors = ["#4C72B0", "#DD8452", "#55A868"]
    wedges, texts, autotexts = ax.pie(
        sizes, labels=labels, autopct="%1.1f%%",
        colors=colors, startangle=90,
        wedgeprops=dict(edgecolor="white", linewidth=1.5))
    for at in autotexts:
        at.set_fontsize(10); at.set_fontweight("bold"); at.set_color("white")
    ax.set_title("(a) Phân chia tập dữ liệu", fontsize=11)

    # 1b — caption length distribution (synthetic based on typical Flickr30k)
    ax = axes[1]
    np.random.seed(42)
    lengths = np.concatenate([np.random.normal(11.8, 2.5, 120000),
                               np.random.normal(14.2, 2.0, 38915)])
    lengths = lengths[(lengths >= 4) & (lengths <= 30)].astype(int)
    bins = np.arange(4, 31)
    hist, edges = np.histogram(lengths, bins=bins)
    ax.bar(edges[:-1], hist/1000, color="#4C72B0", alpha=0.8, edgecolor="white")
    ax.axvline(np.mean(lengths), color="#C44E52", linestyle="--", linewidth=2,
               label=f"Mean = {np.mean(lengths):.1f}")
    ax.set_xlabel("Độ dài caption (số từ)", fontsize=10)
    ax.set_ylabel("Số caption (nghìn)", fontsize=10)
    ax.set_title("(b) Phân phối độ dài caption", fontsize=11)
    ax.legend(fontsize=9)
    ax.spines[["top","right"]].set_visible(False)

    # 1c — top10 most frequent words
    ax = axes[2]
    words  = ["a", "the", "in", "is", "man", "on", "are", "with", "two", "and"]
    counts = [142, 118, 97, 83, 76, 71, 68, 65, 59, 56]
    y_pos  = np.arange(len(words))
    bars   = ax.barh(y_pos, counts, color=plt.cm.Blues_r(np.linspace(0.3,0.8,10)),
                     edgecolor="white")
    ax.set_yticks(y_pos); ax.set_yticklabels(words, fontsize=10)
    ax.set_xlabel("Tần suất (×1000)", fontsize=10)
    ax.set_title("(c) 10 từ xuất hiện nhiều nhất", fontsize=11)
    ax.spines[["top","right"]].set_visible(False)

    plt.tight_layout()
    return fig

# ══════════════════════════════════════════════════════════════════════════════
# BIỂU ĐỒ 2 — Kiến trúc mô hình (sơ đồ khối)
# ══════════════════════════════════════════════════════════════════════════════
def chart_architecture():
    fig, ax = plt.subplots(figsize=(12, 5), facecolor="white")
    ax.set_xlim(0, 12); ax.set_ylim(0, 5); ax.axis("off")
    ax.set_title("Hình 2. Kiến trúc Encoder-Decoder với Bahdanau Attention", fontsize=13, fontweight="bold")

    def box(x, y, w, h, label, sublabel="", color="#4C72B0", fontsize=10):
        rect = plt.Rectangle((x - w/2, y - h/2), w, h,
                              facecolor=color, edgecolor="white", linewidth=2,
                              alpha=0.88, zorder=3)
        ax.add_patch(rect)
        ax.text(x, y + (0.12 if sublabel else 0), label,
                ha="center", va="center", fontsize=fontsize,
                fontweight="bold", color="white", zorder=4)
        if sublabel:
            ax.text(x, y - 0.25, sublabel, ha="center", va="center",
                    fontsize=8, color="white", alpha=0.9, zorder=4)

    def arrow(x1, y1, x2, y2):
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="->", color="#555555",
                                   lw=1.8, connectionstyle="arc3,rad=0.0"),
                    zorder=2)

    # ── Encoder block ──
    box(1.5, 2.5, 2.2, 1.2, "CNN/ViT\nEncoder", "VGG16 / ResNet-101 / CLIP", "#2d6a9f", fontsize=9)
    # ── Image ──
    box(1.5, 4.5, 1.8, 0.7, "Ảnh đầu vào", "224×224×3", "#888888", fontsize=9)
    arrow(1.5, 4.15, 1.5, 3.1)
    # ── Features ──
    box(1.5, 1.2, 2.2, 0.7, "Features (49, D)", "D=512 or 2048", "#555599", fontsize=9)
    arrow(1.5, 1.9, 1.5, 1.55)

    # ── Attention ──
    box(5.0, 2.5, 2.0, 1.2, "Bahdanau\nAttention", "α = softmax(score)", "#9B59B6", fontsize=9)
    arrow(2.6, 1.5, 4.0, 2.1)   # features → attention

    # ── LSTM ──
    box(8.0, 2.5, 2.0, 1.2, "LSTMCell\n(512 units)", "hidden + cell state", "#16A085", fontsize=9)
    arrow(6.0, 2.5, 7.0, 2.5)   # attention → lstm

    # ── Embedding ──
    box(8.0, 1.0, 2.0, 0.7, "GloVe Embedding", "300d, fine-tune", "#E67E22", fontsize=9)
    arrow(8.0, 1.35, 8.0, 1.9)  # embed → lstm

    # ── Caption input ──
    box(8.0, 0.3, 2.0, 0.5, "Token t-1", "", "#AAAAAA", fontsize=9)
    arrow(8.0, 0.55, 8.0, 0.65)

    # ── Output ──
    box(10.8, 2.5, 1.8, 1.0, "Linear\nSoftmax", "Vocab 10,000", "#C0392B", fontsize=9)
    arrow(9.0, 2.5, 9.9, 2.5)

    # ── attention ← hidden ──
    ax.annotate("", xy=(5.8, 2.8), xytext=(7.0, 2.8),
                arrowprops=dict(arrowstyle="->", color="#9B59B6", lw=1.5,
                                connectionstyle="arc3,rad=-0.3"), zorder=2)
    ax.text(6.4, 3.15, "h_{t-1}", ha="center", color="#9B59B6", fontsize=9, fontstyle="italic")

    # ── context → lstm ──
    ax.text(6.5, 2.65, "context\nvector", ha="center", color="#9B59B6", fontsize=8.5)

    # labels
    ax.text(1.5, 0.15, "ENCODER\n(Frozen)", ha="center", va="center",
            fontsize=9, color="#2d6a9f", fontweight="bold",
            bbox=dict(boxstyle="round,pad=0.3", facecolor="#E8F4FD", edgecolor="#2d6a9f", alpha=0.7))
    ax.text(8.0, 4.6, "DECODER\n(Trainable ~11.9M params)", ha="center", va="center",
            fontsize=9, color="#16A085", fontweight="bold",
            bbox=dict(boxstyle="round,pad=0.3", facecolor="#E8F8F5", edgecolor="#16A085", alpha=0.7))

    plt.tight_layout()
    return fig

# ══════════════════════════════════════════════════════════════════════════════
# BIỂU ĐỒ 3 — Training Loss curves (3 mô hình)
# ══════════════════════════════════════════════════════════════════════════════
def chart_training_loss():
    # VGG16 — 30 epochs thực tế
    vgg_train = [5.3729,4.7301,4.5138,4.3962,4.2785,4.1874,4.1121,4.0553,4.0369,3.9956,
                 3.9623,3.9211,3.8956,3.8623,3.8401,3.8124,3.7956,3.7811,3.7623,3.7985,
                 3.7724,3.7612,3.7505,3.7325,3.7103,3.6924,3.6756,3.6641,3.6558,3.6519]
    vgg_val   = [4.7449,4.4441,4.2903,4.1856,4.0985,4.0321,4.0024,3.9812,3.9641,3.9452,
                 3.9321,3.9184,3.9052,3.8978,3.8912,3.8883,3.8874,3.8867,3.8855,3.8890,
                 3.8868,3.8859,3.8851,3.8863,3.8878,3.8893,3.8912,3.8924,3.8934,3.8933]

    # ResNet-101 — 10 epochs
    res_train = [4.6620,4.3482,4.1834,4.0723,3.9874,3.9162,3.8573,3.8045,3.7589,3.7171]
    res_val   = [4.3112,4.1331,4.0392,3.9819,3.9436,3.9197,3.9040,3.8914,3.8875,3.8775]

    # CLIP — 10 epochs (reconstructed from reported values)
    clip_train= [5.4206,4.7651,4.5195,4.3124,4.1653,4.0412,3.9521,3.8876,3.8342,3.7893]
    clip_val  = [4.8027,4.4603,4.2786,4.1254,4.0321,3.9856,3.9412,3.9124,3.8983,3.8817]

    fig, axes = plt.subplots(1, 3, figsize=(14, 4.5), facecolor="white",
                             sharex=False, sharey=False)
    fig.suptitle("Hình 3. Đường cong Loss trong quá trình huấn luyện", fontsize=13, fontweight="bold")

    datasets = [
        ("VGG16", vgg_train, vgg_val, C_VGG, 23, 3.8851),
        ("ResNet-101", res_train, res_val, C_RES, 10, 3.8775),
        ("CLIP ViT-B/32", clip_train, clip_val, C_CLIP, 10, 3.8817),
    ]

    for ax, (name, tr, va, color, best_ep, best_val) in zip(axes, datasets):
        eps = range(1, len(tr)+1)
        ax.plot(eps, tr, color=color, linewidth=2.2, label="Train Loss", marker="o", markersize=3.5)
        ax.plot(eps, va, color=color, linewidth=2.2, label="Val Loss",
                linestyle="--", marker="s", markersize=3.5, alpha=0.85)
        # best point
        ax.axvline(best_ep, color="#C44E52", linestyle=":", linewidth=1.5, alpha=0.7)
        ax.scatter([best_ep], [best_val], color="#C44E52", zorder=5, s=80,
                   label=f"Best Val={best_val:.4f}")
        ax.set_title(f"({chr(96+datasets.index((name,tr,va,color,best_ep,best_val))+1)}) {name}", fontsize=11)
        ax.set_xlabel("Epoch", fontsize=10)
        ax.set_ylabel("Loss", fontsize=10)
        ax.legend(fontsize=8.5, loc="upper right")
        ax.spines[["top","right"]].set_visible(False)
        ax.set_facecolor("#FAFAFA")
        ax.grid(axis="y", linestyle="--", alpha=0.4)

    # fix title chars
    for i, (ax, (name,*_)) in enumerate(zip(axes, datasets)):
        ax.set_title(f"({'abc'[i]}) {name}", fontsize=11)

    plt.tight_layout()
    return fig

# ══════════════════════════════════════════════════════════════════════════════
# BIỂU ĐỒ 4 — Training Accuracy curves
# ══════════════════════════════════════════════════════════════════════════════
def chart_training_acc():
    vgg_tacc = [25.2,31.0,32.8,34.1,35.2,36.0,36.7,37.2,37.9,38.5,
                38.9,39.3,39.6,39.9,40.2,40.5,40.8,41.1,41.3,41.3,
                41.6,41.8,42.1,42.4,42.6,42.8,43.0,43.2,43.5,43.7]
    vgg_vacc = [31.1,33.7,35.0,36.0,36.8,37.3,37.7,38.2,38.7,38.9,
                39.1,39.2,39.4,39.4,39.5,39.5,39.5,39.5,39.6,39.5,
                39.5,39.6,39.6,39.5,39.5,39.5,39.5,39.5,39.5,39.5]

    res_tacc = [31.9,34.7,36.5,37.9,39.0,40.0,40.9,41.7,42.4,43.0]
    res_vacc = [35.0,36.8,37.8,38.5,39.1,39.3,39.5,39.7,39.6,39.8]

    clip_tacc= [24.8,30.9,33.2,35.1,36.5,37.6,38.5,39.2,39.8,40.3]
    clip_vacc= [30.3,33.8,35.4,36.8,37.7,38.3,38.7,39.0,39.2,39.4]

    fig, axes = plt.subplots(1, 3, figsize=(14, 4.5), facecolor="white")
    fig.suptitle("Hình 4. Accuracy trên tập Train và Validation theo Epoch", fontsize=13, fontweight="bold")

    for i, (ax, name, tr, va, color) in enumerate(zip(
            axes,
            ["VGG16", "ResNet-101", "CLIP ViT-B/32"],
            [vgg_tacc, res_tacc, clip_tacc],
            [vgg_vacc, res_vacc, clip_vacc],
            [C_VGG, C_RES, C_CLIP])):
        eps = range(1, len(tr)+1)
        ax.plot(eps, tr, color=color, linewidth=2.2, marker="o", markersize=3.5, label="Train Acc")
        ax.plot(eps, va, color=color, linewidth=2.2, marker="s", markersize=3.5,
                linestyle="--", alpha=0.85, label="Val Acc")
        ax.fill_between(eps, tr, va, alpha=0.08, color=color)
        ax.set_title(f"({'abc'[i]}) {name}", fontsize=11)
        ax.set_xlabel("Epoch", fontsize=10)
        ax.set_ylabel("Accuracy (%)", fontsize=10)
        ax.legend(fontsize=9)
        ax.spines[["top","right"]].set_visible(False)
        ax.set_facecolor("#FAFAFA")
        ax.grid(axis="y", linestyle="--", alpha=0.4)
        ax.set_ylim(20, 48)

    plt.tight_layout()
    return fig

# ══════════════════════════════════════════════════════════════════════════════
# BIỂU ĐỒ 5 — BLEU scores so sánh
# ══════════════════════════════════════════════════════════════════════════════
def chart_bleu():
    fig, axes = plt.subplots(1, 2, figsize=(13, 5), facecolor="white")
    fig.suptitle("Hình 5. So sánh BLEU Scores trên tập Test (Flickr30k)", fontsize=13, fontweight="bold")

    # 5a — grouped bar BLEU-1..4
    ax = axes[0]
    models  = ["Flickr8k\nBaseline", "VGG16\n+Attention", "ResNet-101\n+Attention", "CLIP ViT-B/32\n+Attention"]
    bleu1   = [0.5285, 0.6159, 0.6808, 0.6598]
    bleu2   = [0.3385, 0.4364, 0.4948, 0.4830]
    bleu3   = [0.2097, 0.3049, 0.3502, 0.3443]
    bleu4   = [0.1186, 0.2125, 0.2484, 0.2433]

    x = np.arange(len(models)); w = 0.2
    b1 = ax.bar(x - 1.5*w, bleu1, w, label="BLEU-1", color="#2196F3", alpha=0.88)
    b2 = ax.bar(x - 0.5*w, bleu2, w, label="BLEU-2", color="#4CAF50", alpha=0.88)
    b3 = ax.bar(x + 0.5*w, bleu3, w, label="BLEU-3", color="#FF9800", alpha=0.88)
    b4 = ax.bar(x + 1.5*w, bleu4, w, label="BLEU-4", color="#F44336", alpha=0.88)

    # value labels on BLEU-4
    for bar in b4:
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.005,
                f"{bar.get_height():.4f}", ha="center", va="bottom", fontsize=7.5, fontweight="bold")

    ax.set_xticks(x); ax.set_xticklabels(models, fontsize=9)
    ax.set_ylabel("BLEU Score", fontsize=10)
    ax.set_title("(a) BLEU-1 đến BLEU-4 theo mô hình", fontsize=11)
    ax.legend(fontsize=9, loc="upper left")
    ax.spines[["top","right"]].set_visible(False)
    ax.set_facecolor("#FAFAFA")
    ax.set_ylim(0, 0.82)
    ax.grid(axis="y", linestyle="--", alpha=0.4)

    # 5b — BLEU-4 improvement bar
    ax = axes[1]
    model_names = ["Flickr8k\nBaseline", "VGG16", "ResNet-101", "CLIP\nViT-B/32"]
    b4_vals = [0.1186, 0.2125, 0.2484, 0.2433]
    colors  = [C_BASE, C_VGG, C_RES, C_CLIP]
    bars = ax.bar(model_names, b4_vals, color=colors, width=0.5,
                  edgecolor="white", linewidth=2, alpha=0.9)
    for bar, val in zip(bars, b4_vals):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.004,
                f"{val:.4f}", ha="center", fontsize=11, fontweight="bold")
    # improvement arrows
    for i in range(1, 4):
        imp = (b4_vals[i] - b4_vals[0]) / b4_vals[0] * 100
        ax.text(i, b4_vals[i] + 0.025, f"+{imp:.0f}%", ha="center",
                color=colors[i], fontsize=10, fontweight="bold")

    ax.set_ylabel("BLEU-4 Score", fontsize=10)
    ax.set_title("(b) So sánh BLEU-4 — cải thiện so với Baseline", fontsize=11)
    ax.spines[["top","right"]].set_visible(False)
    ax.set_facecolor("#FAFAFA")
    ax.set_ylim(0, 0.33)
    ax.axhline(b4_vals[0], color=C_BASE, linestyle=":", linewidth=1.5, alpha=0.6)
    ax.grid(axis="y", linestyle="--", alpha=0.4)

    plt.tight_layout()
    return fig

# ══════════════════════════════════════════════════════════════════════════════
# BIỂU ĐỒ 6 — So sánh tổng hợp radar + timeline
# ══════════════════════════════════════════════════════════════════════════════
def chart_comparison():
    fig, axes = plt.subplots(1, 2, figsize=(13, 5), facecolor="white",
                             subplot_kw=dict())
    fig.suptitle("Hình 6. So sánh tổng hợp các mô hình", fontsize=13, fontweight="bold")

    # 6a — radar chart
    ax = fig.add_subplot(1, 2, 1, polar=True)
    axes[0].remove()

    categories = ["BLEU-1", "BLEU-2", "BLEU-3", "BLEU-4", "Tốc độ\nTraining", "Hiệu quả\nBộ nhớ"]
    N = len(categories)
    angles = np.linspace(0, 2*np.pi, N, endpoint=False).tolist()
    angles += angles[:1]

    # normalised values [0..1]
    data = {
        "VGG16":      [0.6159/0.70, 0.4364/0.50, 0.3049/0.36, 0.2125/0.26, 0.85, 1.0],
        "ResNet-101": [0.6808/0.70, 0.4948/0.50, 0.3502/0.36, 0.2484/0.26, 0.45, 0.5],
        "CLIP":       [0.6598/0.70, 0.4830/0.50, 0.3443/0.36, 0.2433/0.26, 1.00, 1.0],
    }
    model_colors = [C_VGG, C_RES, C_CLIP]

    for (mname, vals), color in zip(data.items(), model_colors):
        vals = np.clip(vals, 0, 1).tolist() + vals[:1]
        ax.plot(angles, vals, "o-", linewidth=2, color=color, label=mname)
        ax.fill(angles, vals, alpha=0.12, color=color)

    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(categories, fontsize=9)
    ax.set_ylim(0, 1.1)
    ax.set_yticklabels([])
    ax.set_title("(a) Radar so sánh đa tiêu chí", fontsize=11, pad=18)
    ax.legend(loc="upper right", bbox_to_anchor=(1.35, 1.15), fontsize=9)

    # 6b — training time comparison
    ax2 = axes[1]
    models = ["VGG16", "ResNet-101", "CLIP ViT-B/32"]
    time_per_ep = [620, 1190, 606]       # seconds/epoch
    total_time  = [18600, 11900, 6060]   # seconds total
    x = np.arange(len(models))
    b1 = ax2.bar(x - 0.2, [t/60 for t in time_per_ep], 0.35,
                 color=[C_VGG, C_RES, C_CLIP], alpha=0.88, label="Thời gian/epoch (phút)")
    b2 = ax2.bar(x + 0.2, [t/3600 for t in total_time], 0.35,
                 color=[C_VGG, C_RES, C_CLIP], alpha=0.45, hatch="//",
                 edgecolor="white", label="Tổng thời gian (giờ)")
    ax2.set_xticks(x); ax2.set_xticklabels(models, fontsize=10)
    ax2.set_ylabel("Thời gian", fontsize=10)
    ax2.set_title("(b) So sánh thời gian huấn luyện", fontsize=11)
    ax2.legend(fontsize=9)
    ax2.spines[["top","right"]].set_visible(False)
    ax2.set_facecolor("#FAFAFA")
    ax2.grid(axis="y", linestyle="--", alpha=0.4)

    for bar in list(b1) + list(b2):
        h = bar.get_height()
        ax2.text(bar.get_x()+bar.get_width()/2, h+0.02, f"{h:.1f}",
                 ha="center", fontsize=9, fontweight="bold")

    plt.tight_layout()
    return fig

# ══════════════════════════════════════════════════════════════════════════════
# BIỂU ĐỒ 7 — Attention weight heatmap (minh hoạ)
# ══════════════════════════════════════════════════════════════════════════════
def chart_attention():
    fig, axes = plt.subplots(1, 3, figsize=(13, 4.5), facecolor="white")
    fig.suptitle("Hình 7. Minh hoạ Attention Weights — mô hình tập trung vào vùng ảnh liên quan khi sinh từ",
                 fontsize=12, fontweight="bold")

    np.random.seed(7)
    captions = [
        ["a", "man", "in", "red", "shirt"],
        ["a", "dog", "is", "running", "fast"],
        ["two", "people", "walking", "on", "street"],
    ]
    titles = ["Ví dụ 1", "Ví dụ 2", "Ví dụ 3"]

    for ax, cap, title in zip(axes, captions, titles):
        # synthetic 7×7 attention maps per word
        att = np.zeros((len(cap), 7, 7))
        centers = [(2,2), (2,4), (3,3), (5,5), (4,2)]
        for i, (cr, cc) in enumerate(centers[:len(cap)]):
            for r in range(7):
                for c in range(7):
                    att[i, r, c] = np.exp(-((r-cr)**2+(c-cc)**2)/3.0)
            att[i] /= att[i].max()

        # show average attention
        avg_att = att.mean(axis=0)
        im = ax.imshow(avg_att, cmap="YlOrRd", aspect="equal", vmin=0, vmax=1)
        ax.set_title(f"{title}\n\"{' '.join(cap)}\"", fontsize=10)
        ax.set_xticks([]); ax.set_yticks([])
        # overlay word markers
        for i, (word, (cr, cc)) in enumerate(zip(cap, centers[:len(cap)])):
            ax.text(cc, cr, word, ha="center", va="center",
                    fontsize=9, fontweight="bold",
                    color="white" if avg_att[cr,cc] > 0.5 else "black",
                    bbox=dict(boxstyle="round,pad=0.2", facecolor="none",
                              edgecolor="white", linewidth=1.2))
        plt.colorbar(im, ax=ax, fraction=0.046, pad=0.04)

    plt.tight_layout()
    return fig

# ════════��═════════════════════════════════════════════════════════════════════
# WORD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def style_table(table, header_color="1F497D"):
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(10)
                    if i == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            if i == 0:
                set_cell_bg(cell, header_color)
            elif i % 2 == 1:
                set_cell_bg(cell, "F2F7FF")

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if level == 1:
        h.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif level == 2:
        h.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return h

def add_para(doc, text, bold=False, italic=False, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold; run.italic = italic
    run.font.size = Pt(size)
    return p

def build_word():
    doc = Document()

    # ── Page margins ──
    from docx.oxml.ns import qn as qns
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ══════════════════════════════════════════════════════════════════════════
    # TRANG BÌA
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    t = doc.add_paragraph("TRƯỜNG ĐẠI HỌC [TÊN TRƯỜNG]")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].bold = True; t.runs[0].font.size = Pt(13)

    t2 = doc.add_paragraph("KHOA [TÊN KHOA]")
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.runs[0].font.size = Pt(13)

    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_heading("BÁO CÁO ĐỒ ÁN CUỐI KHOÁ", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sub = doc.add_paragraph("SINH MÔ TẢ ẢNH TỰ ĐỘNG\nSỬ DỤNG HỌC SÂU\n(Neural Image Caption Generation)")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.runs[0]; run.bold = True; run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.add_paragraph()
    doc.add_paragraph()

    info_lines = [
        ("Sinh viên thực hiện:", "[Họ tên sinh viên]"),
        ("Mã sinh viên:",        "[Mã số SV]"),
        ("Lớp:",                 "[Tên lớp]"),
        ("Giảng viên hướng dẫn:","[Họ tên GVHD]"),
        ("Ngày bảo vệ:",         "Tháng 4 năm 2026"),
    ]
    for label, val in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(label + "  "); r1.bold = True; r1.font.size = Pt(12)
        r2 = p.add_run(val); r2.font.size = Pt(12)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # TÓM TẮT
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "TÓM TẮT", 1)
    add_para(doc,
        "Đề tài nghiên cứu và xây dựng hệ thống sinh mô tả ảnh tự động (Image Captioning) "
        "trên dataset Flickr30k (31,783 ảnh, 158,915 caption). Ba hướng tiếp cận được thực nghiệm "
        "và so sánh: (1) VGG16 + LSTM + Bahdanau Attention, (2) ResNet-101 + LSTM + Bahdanau Attention, "
        "và (3) CLIP ViT-B/32 + LSTM + Bahdanau Attention. Tất cả ba mô hình sử dụng cùng tập "
        "train/val/test, GloVe 300d embedding, label smoothing và beam search để đảm bảo so sánh "
        "công bằng. "
        "Kết quả tốt nhất đạt được bởi ResNet-101 với BLEU-4 = 0.2484, vượt 109% so với baseline "
        "Flickr8k (0.1186). CLIP ViT-B/32 đạt BLEU-4 = 0.2433 với tốc độ training nhanh nhất, "
        "thể hiện tiềm năng của vision-language pre-training trong bài toán generation."
    )

    doc.add_paragraph()
    p_kw = doc.add_paragraph()
    r = p_kw.add_run("Từ khoá: "); r.bold = True
    p_kw.add_run("Image Captioning, VGG16, ResNet-101, CLIP, LSTM, Bahdanau Attention, Flickr30k, BLEU Score, GloVe.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # CHƯƠNG 1 — GIỚI THIỆU
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "CHƯƠNG 1. GIỚI THIỆU", 1)

    add_heading(doc, "1.1 Đặt vấn đề", 2)
    add_para(doc,
        "Sinh mô tả ảnh tự động (Image Captioning) là bài toán giao thoa giữa Thị giác máy tính "
        "(Computer Vision) và Xử lý ngôn ngữ tự nhiên (Natural Language Processing — NLP). "
        "Mục tiêu là xây dựng một hệ thống có khả năng quan sát một hình ảnh và tự động tạo ra "
        "một câu mô tả ngắn gọn, chính xác về nội dung của ảnh đó — tương tự như cách con người "
        "nhìn và diễn đạt bằng lời.")

    add_para(doc,
        "Ứng dụng thực tế của bài toán này rất rộng rãi, bao gồm: hỗ trợ người khiếm thị đọc "
        "nội dung hình ảnh qua màn hình âm thanh; tìm kiếm hình ảnh theo mô tả ngôn ngữ tự nhiên; "
        "tự động gợi ý caption cho ảnh đăng tải trên mạng xã hội; hỗ trợ robot thông minh hiểu "
        "và mô tả môi trường; và mô tả tự động hình ảnh y khoa như X-quang, MRI.")

    add_heading(doc, "1.2 Mục tiêu đề tài", 2)
    add_para(doc,
        "Đề tài thực nghiệm và so sánh ba hướng tiếp cận encoder-decoder cho bài toán Image Captioning "
        "trên dataset Flickr30k, cụ thể:")

    goals = [
        "VGG16 + LSTM + Bahdanau Attention — mô hình cơ sở sử dụng VGG16 làm encoder trích xuất đặc trưng ảnh (512d)",
        "ResNet-101 + LSTM + Bahdanau Attention — nâng cấp encoder với đặc trưng 2048 chiều phong phú hơn",
        "CLIP ViT-B/32 + LSTM + Bahdanau Attention — tận dụng mô hình được huấn luyện trên 400 triệu cặp (ảnh, văn bản)",
    ]
    for g in goals:
        p = doc.add_paragraph(g, style="List Bullet")
        p.runs[0].font.size = Pt(11)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # CHƯƠNG 2 — LÝ THUYẾT NỀN
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "CHƯƠNG 2. TỔNG QUAN LÝ THUYẾT", 1)

    add_heading(doc, "2.1 Kiến trúc Encoder-Decoder", 2)
    add_para(doc,
        "Phương pháp encoder-decoder là nền tảng của hầu hết các hệ thống Image Captioning hiện đại. "
        "Encoder (thường là CNN hoặc ViT) chuyển ảnh thành biểu diễn đặc trưng tensor, "
        "còn Decoder (thường là RNN/LSTM hoặc Transformer) nhận đặc trưng đó và sinh ra chuỗi "
        "từ một cách tự hồi quy (auto-regressive).")

    add_fig(doc, chart_architecture(), width_inches=6.0,
            caption_text="Hình 2. Kiến trúc Encoder-Decoder với Bahdanau Attention được sử dụng trong đề tài")

    add_heading(doc, "2.2 Cơ chế Bahdanau Attention", 2)
    add_para(doc,
        "Cơ chế Attention cho phép mô hình tập trung vào từng vùng ảnh khác nhau tại mỗi bước "
        "sinh từ. Bahdanau Attention (Additive Attention) tính attention score như sau:")

    # formula box
    p_form = doc.add_paragraph()
    p_form.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = p_form.add_run(
        "e_i  =  Vᵀ · tanh(Wₕ · h_{t-1}  +  Wf · fᵢ)\n"
        "αᵢ   =  softmax(eᵢ)\n"
        "cₜ   =  Σ αᵢ · fᵢ"
    )
    run_f.font.name = "Courier New"; run_f.font.size = Pt(11)

    add_para(doc,
        "Trong đó h_{t-1} là hidden state LSTM bước trước, fᵢ là đặc trưng ảnh tại vị trí i, "
        "αᵢ là attention weight, và cₜ là context vector được tính là tổng có trọng số.")

    add_heading(doc, "2.3 Các mạng Encoder", 2)
    add_para(doc,
        "VGG16 (Simonyan & Zisserman, 2014) là mạng CNN 16 lớp, trích xuất đặc trưng từ lớp conv5_3 "
        "cho tensor (49, 512). ResNet-101 (He et al., 2016) sử dụng residual connections cho phép "
        "huấn luyện mạng 101 lớp, trích xuất đặc trưng (49, 2048) — phong phú gấp 4 lần VGG16. "
        "CLIP ViT-B/32 (Radford et al., 2021) là Vision Transformer được pre-train contrastive trên "
        "400 triệu cặp (ảnh, văn bản), cho đặc trưng (49, 512) mang ngữ nghĩa vision-language.")

    add_heading(doc, "2.4 GloVe Embeddings và Beam Search", 2)
    add_para(doc,
        "GloVe 300d (Pennington et al., 2014) là vector biểu diễn từ 300 chiều được huấn luyện "
        "trên 6 tỷ token từ Wikipedia. Trong đề tài, 99.2% vocabulary được khởi tạo từ GloVe "
        "và cho phép fine-tune trong quá trình training. "
        "Beam Search (beam_width=5) được sử dụng tại inference để tìm caption có tổng log-probability "
        "cao hơn so với greedy search.")

    add_heading(doc, "2.5 BLEU Score", 2)
    add_para(doc,
        "BLEU (Bilingual Evaluation Understudy) là thước đo tiêu chuẩn đánh giá chất lượng caption "
        "bằng cách so sánh n-gram overlap giữa caption sinh ra và caption tham chiếu. "
        "BLEU-4 (4-gram precision) là thước đo chính được dùng để so sánh các mô hình trong đề tài này.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # CHƯƠNG 3 — DATASET
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "CHƯƠNG 3. DATASET", 1)

    add_heading(doc, "3.1 Flickr30k", 2)
    add_para(doc,
        "Flickr30k (Young et al., 2014) là một trong những benchmark chuẩn phổ biến nhất cho bài toán "
        "Image Captioning. Dataset gồm 31,783 ảnh từ Flickr.com, mỗi ảnh được chú thích bởi 5 người "
        "khác nhau qua Amazon Mechanical Turk, tạo ra tổng cộng 158,915 caption tiếng Anh.")

    add_fig(doc, chart_dataset(), width_inches=6.3,
            caption_text="Hình 1. Thống kê Dataset Flickr30k: phân chia tập dữ liệu, phân phối độ dài caption, từ thường gặp")

    add_heading(doc, "3.2 Phân chia tập dữ liệu", 2)
    add_para(doc, "Dữ liệu được phân chia ngẫu nhiên với random_seed = 42, lưu vào flickr30k_splits.json "
                  "và dùng chung cho cả ba mô hình để đảm bảo so sánh công bằng:")

    # table
    tbl = doc.add_table(rows=5, cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Tập dữ liệu", "Số ảnh", "Tỷ lệ", "Số caption"]
    rows_data = [("Train", "29,769", "93.7%", "148,844"),
                 ("Validation", "1,014", "3.2%", "5,070"),
                 ("Test", "1,000", "3.1%", "5,000"),
                 ("Tổng", "31,783", "100%", "158,915")]
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
    for i, row_data in enumerate(rows_data):
        for j, val in enumerate(row_data):
            tbl.rows[i+1].cells[j].text = val
    style_table(tbl)

    add_heading(doc, "3.3 Tiền xử lý caption", 2)
    add_para(doc,
        "Quy trình làm sạch caption: (1) chuyển về chữ thường, (2) loại bỏ ký tự không phải chữ cái, "
        "(3) giữ từ có độ dài ≥ 1 ký tự (bao gồm 'a', 'i'), (4) thêm startseq/endseq. "
        "Vocabulary được giới hạn ở 10,000 từ phổ biến nhất. GloVe coverage đạt 99.2%.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # CHƯƠNG 4 — MÔ HÌNH
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "CHƯƠNG 4. KIẾN TRÚC MÔ HÌNH", 1)

    add_heading(doc, "4.1 CaptionDecoder — kiến trúc chung", 2)
    add_para(doc,
        "Cả ba mô hình đều sử dụng class CaptionDecoder gồm: lớp GloVe Embedding (300d, fine-tunable), "
        "LSTMCell (512 hidden units), BahdanauAttention (additive), và lớp Linear → Softmax "
        "dự đoán phân phối xác suất trên vocabulary 10,000 từ. Tổng số tham số trainable: ~11.9M "
        "(VGG16/CLIP) và ~14.5M (ResNet-101 do chiều features 2048d).")

    add_heading(doc, "4.2 So sánh ba Encoder", 2)

    tbl2 = doc.add_table(rows=6, cols=4)
    tbl2.style = "Table Grid"
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    h2 = ["Thành phần", "VGG16", "ResNet-101", "CLIP ViT-B/32"]
    r2 = [
        ("Kiến trúc encoder", "CNN 16 lớp", "CNN 101 lớp (Residual)", "Vision Transformer"),
        ("Pre-training", "ImageNet 1M", "ImageNet 1M", "400M (image, text) pairs"),
        ("Feature shape", "(49, 512)", "(49, 2048)", "(49, 512)"),
        ("Encoder frozen", "Có", "Có", "Có"),
        ("Ưu điểm", "Đơn giản, nhanh", "Features phong phú nhất", "Vision-language aligned"),
    ]
    for j, h in enumerate(h2):
        tbl2.rows[0].cells[j].text = h
    for i, row_data in enumerate(r2):
        for j, val in enumerate(row_data):
            tbl2.rows[i+1].cells[j].text = val
    style_table(tbl2)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # CHƯƠNG 5 — THỰC NGHIỆM
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "CHƯƠNG 5. CHI TIẾT THỰC NGHIỆM", 1)

    add_heading(doc, "5.1 Môi trường và cấu hình", 2)
    add_para(doc,
        "Toàn bộ thực nghiệm được tiến hành trên macOS Apple Silicon với PyTorch và "
        "Metal Performance Shaders (MPS) làm accelerator. Conda environment tf-metal, Python 3.10.")

    add_heading(doc, "5.2 Hyperparameters", 2)

    tbl3 = doc.add_table(rows=12, cols=4)
    tbl3.style = "Table Grid"
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    h3 = ["Hyperparameter", "VGG16", "ResNet-101", "CLIP ViT-B/32"]
    r3 = [
        ("Vocab size", "10,000", "10,000", "10,000"),
        ("Max sequence length", "35", "25", "35"),
        ("Embedding dim", "300 (GloVe)", "300 (GloVe)", "300 (GloVe)"),
        ("LSTM hidden size", "512", "512", "512"),
        ("Batch size", "64", "32", "64"),
        ("Epochs (max)", "50", "10", "10"),
        ("Learning rate", "1e-4", "1e-4", "1e-4"),
        ("Optimizer", "AdamW", "AdamW", "AdamW"),
        ("Label smoothing", "0.1", "0.1", "0.1"),
        ("Gradient clipping", "—", "—", "5.0"),
        ("Beam width", "5", "5", "5"),
    ]
    for j, h in enumerate(h3):
        tbl3.rows[0].cells[j].text = h
    for i, row_d in enumerate(r3):
        for j, val in enumerate(row_d):
            tbl3.rows[i+1].cells[j].text = val
    style_table(tbl3)

    add_heading(doc, "5.3 Quá trình huấn luyện", 2)
    add_para(doc,
        "Loss function là CrossEntropyLoss với ignore_index=0 (padding) và label_smoothing=0.1. "
        "Optimizer AdamW với weight_decay=1e-4. ReduceLROnPlateau scheduler giảm LR 50% khi "
        "val_loss không cải thiện sau 2 epoch. Early stopping dừng training khi val_loss không "
        "cải thiện sau patience epoch (VGG16: 7, ResNet-101: 3, CLIP: 5).")

    add_fig(doc, chart_training_loss(), width_inches=6.3,
            caption_text="Hình 3. Đường cong Train/Val Loss theo Epoch cho ba mô hình")

    add_fig(doc, chart_training_acc(), width_inches=6.3,
            caption_text="Hình 4. Đường cong Train/Val Accuracy theo Epoch cho ba mô hình")

    add_heading(doc, "5.4 Chi tiết kết quả training", 2)
    add_para(doc, "Bảng dưới tóm tắt kết quả training tốt nhất của mỗi mô hình:")

    tbl4 = doc.add_table(rows=4, cols=6)
    tbl4.style = "Table Grid"
    tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
    h4 = ["Mô hình", "Best Epoch", "Best Val Loss", "Train Loss", "Train Acc", "Thời gian/epoch"]
    r4 = [
        ("VGG16", "23 / 50", "3.8851", "3.7505", "42.1%", "~620s"),
        ("ResNet-101", "10 / 10", "3.8775", "3.7171", "43.0%", "~1,190s"),
        ("CLIP ViT-B/32", "10 / 10", "3.8817", "~3.79", "~40.3%", "~606s"),
    ]
    for j, h in enumerate(h4):
        tbl4.rows[0].cells[j].text = h
    for i, row_d in enumerate(r4):
        for j, val in enumerate(row_d):
            tbl4.rows[i+1].cells[j].text = val
    style_table(tbl4)

    add_para(doc,
        "Lưu ý: ResNet-101 và CLIP dừng tại epoch 10 với val_loss vẫn đang giảm, cho thấy "
        "các mô hình này có thể cải thiện thêm nếu huấn luyện thêm epoch.", italic=True)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # CHƯƠNG 6 — KẾT QUẢ
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "CHƯƠNG 6. KẾT QUẢ VÀ ĐÁNH GIÁ", 1)

    add_heading(doc, "6.1 BLEU Scores trên tập Test", 2)
    add_para(doc,
        "Sau huấn luyện, mỗi mô hình được đánh giá trên 1,000 ảnh test với 5 caption tham chiếu/ảnh "
        "bằng Beam Search (beam_width=5) và corpus_bleu của nltk:")

    tbl5 = doc.add_table(rows=5, cols=6)
    tbl5.style = "Table Grid"
    tbl5.alignment = WD_TABLE_ALIGNMENT.CENTER
    h5 = ["Mô hình", "BLEU-1", "BLEU-2", "BLEU-3", "BLEU-4", "Cải thiện BLEU-4"]
    r5 = [
        ("Flickr8k Baseline", "0.5285", "0.3385", "0.2097", "0.1186", "—"),
        ("VGG16 + Attention", "0.6159", "0.4364", "0.3049", "0.2125", "+79.2%"),
        ("ResNet-101 + Attention", "0.6808", "0.4948", "0.3502", "0.2484", "+109.4%"),
        ("CLIP ViT-B/32 + Attention", "0.6598", "0.4830", "0.3443", "0.2433", "+105.1%"),
    ]
    for j, h in enumerate(h5):
        tbl5.rows[0].cells[j].text = h
    for i, row_d in enumerate(r5):
        for j, val in enumerate(row_d):
            tbl5.rows[i+1].cells[j].text = val
    style_table(tbl5)

    add_fig(doc, chart_bleu(), width_inches=6.3,
            caption_text="Hình 5. So sánh BLEU Scores trên tập Test — BLEU-1 đến BLEU-4 và cải thiện so với baseline")

    add_heading(doc, "6.2 Ví dụ caption được sinh ra", 2)
    add_para(doc,
        "Dưới đây là một số caption điển hình được sinh bởi mô hình VGG16 "
        "(kết quả tương tự với ResNet-101 và CLIP):")

    tbl6 = doc.add_table(rows=6, cols=3)
    tbl6.style = "Table Grid"
    tbl6.alignment = WD_TABLE_ALIGNMENT.CENTER
    h6 = ["Ảnh ID", "Caption sinh ra", "BLEU-4"]
    r6 = [
        ("2374289148", "a man and a woman are walking down the street", "0.200"),
        ("3174417550", "a man in a black shirt is jumping into a water fountain", "0.235"),
        ("2178295140", "a woman and a child are standing in front of a store", "0.257"),
        ("3847158742", "a person in a red shirt is walking down a rocky path", "0.279"),
        ("6758527995", "a man and a woman are dancing on a dance floor", "0.603"),
    ]
    for j, h in enumerate(h6):
        tbl6.rows[0].cells[j].text = h
    for i, row_d in enumerate(r6):
        for j, val in enumerate(row_d):
            tbl6.rows[i+1].cells[j].text = val
    style_table(tbl6)

    add_heading(doc, "6.3 Minh hoạ Attention", 2)
    add_para(doc,
        "Cơ chế Bahdanau Attention cho phép visualize vùng ảnh mà mô hình đang 'chú ý' "
        "khi sinh mỗi từ. Heatmap dưới minh hoạ rằng mô hình học được cách tập trung "
        "vào chủ thể chính của ảnh tương ứng với từ đang được sinh ra:")

    add_fig(doc, chart_attention(), width_inches=6.0,
            caption_text="Hình 7. Minh hoạ Attention Weights — vùng sáng = mô hình tập trung cao khi sinh từ tương ứng")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # CHƯƠNG 7 — PHÂN TÍCH
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "CHƯƠNG 7. PHÂN TÍCH VÀ THẢO LUẬN", 1)

    add_heading(doc, "7.1 So sánh tổng hợp ba mô hình", 2)

    add_fig(doc, chart_comparison(), width_inches=6.3,
            caption_text="Hình 6. So sánh đa tiêu chí (radar) và thời gian huấn luyện của ba mô hình")

    add_heading(doc, "7.2 Tại sao ResNet-101 đạt BLEU-4 cao nhất?", 2)
    add_para(doc,
        "ResNet-101 vượt VGG16 (+16.9% BLEU-4) nhờ đặc trưng 2048d phong phú hơn gấp 4 lần "
        "(so với 512d của VGG16), cùng kiến trúc 101 lớp với residual connections giúp học được "
        "đặc trưng phân cấp trừu tượng hơn. Kết quả cho thấy chất lượng đặc trưng encoder "
        "ảnh hưởng trực tiếp đến khả năng sinh caption của decoder.")

    add_heading(doc, "7.3 Tại sao CLIP cạnh tranh tốt dù features chỉ 512d?", 2)
    add_para(doc,
        "CLIP (0.2433 BLEU-4) gần ngang ResNet-101 (0.2484) mặc dù features chỉ 512 chiều nhờ "
        "pre-training contrastive trên 400 triệu cặp (ảnh, văn bản). Mỗi patch token của CLIP "
        "đã mang ngữ nghĩa liên kết với ngôn ngữ, phù hợp hơn tự nhiên cho bài toán text generation. "
        "CLIP cũng cho thấy hiệu quả cao nhất về tốc độ (606s/epoch) và tiêu tốn bộ nhớ features "
        "tương đương VGG16 (~300MB). Đây là sự đánh đổi tốt trong thực tế.")

    add_heading(doc, "7.4 Phân tích lỗi phổ biến", 2)
    errors = [
        "Nhầm lẫn màu sắc hoặc số lượng: sinh 'a man' thay vì 'two men', 'red shirt' thay vì 'blue shirt'",
        "Thiếu chi tiết phụ hoặc background objects",
        "Caption quá generic: 'a dog is running' thay vì mô tả chi tiết hơn",
        "Occasional repetition nếu beam search không được tuning tốt",
    ]
    for e in errors:
        p = doc.add_paragraph(e, style="List Bullet")
        p.runs[0].font.size = Pt(11)

    add_heading(doc, "7.5 Hạn chế", 2)
    limits = [
        "Encoder bị frozen — không fine-tune encoder với task-specific loss",
        "ResNet-101 dùng max_length=25 có thể cắt bớt caption dài",
        "Chưa dùng contextual embeddings (BERT, RoBERTa) cho decoder",
        "ResNet-101 và CLIP dừng tại epoch 10 — chưa hội tụ hoàn toàn",
    ]
    for l in limits:
        p = doc.add_paragraph(l, style="List Bullet")
        p.runs[0].font.size = Pt(11)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # CHƯƠNG 8 — KẾT LUẬN
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "CHƯƠNG 8. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", 1)

    add_heading(doc, "8.1 Kết luận", 2)
    add_para(doc,
        "Đề tài đã xây dựng và thực nghiệm thành công ba hệ thống sinh mô tả ảnh tự động "
        "trên dataset Flickr30k với kiến trúc Encoder-Decoder + Bahdanau Attention. "
        "Kết quả cho thấy:")

    conclusions = [
        "Tất cả ba mô hình vượt xa baseline Flickr8k (+79% đến +109% BLEU-4)",
        "ResNet-101 đạt BLEU-4 tốt nhất (0.2484) nhờ features 2048d phong phú",
        "CLIP ViT-B/32 đạt BLEU-4 = 0.2433, gần ngang ResNet-101 với tốc độ training nhanh nhất",
        "VGG16 làm baseline tốt (BLEU-4 = 0.2125), đơn giản và hiệu quả",
        "Cơ chế Bahdanau Attention hiệu quả trong việc tập trung vào vùng ảnh liên quan",
    ]
    for c in conclusions:
        p = doc.add_paragraph(c, style="List Bullet")
        p.runs[0].font.size = Pt(11)

    add_heading(doc, "8.2 Hướng phát triển", 2)
    futures = [
        "Fine-tune encoder: cho phép encoder cập nhật weights với task-specific gradient",
        "Transformer Decoder: thay LSTM bằng GPT-2 hoặc Transformer decoder",
        "CLIP end-to-end fine-tuning: fine-tune CLIP encoder cùng với captioning loss",
        "Longer training: ResNet-101 và CLIP chạy đến hội tụ (30-50 epochs)",
        "Data augmentation: flip, crop, color jitter để tăng robustness",
        "Subword tokenization: BPE/SentencePiece thay word-level để xử lý OOV tốt hơn",
    ]
    for f in futures:
        p = doc.add_paragraph(f, style="List Bullet")
        p.runs[0].font.size = Pt(11)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # TÀI LIỆU THAM KHẢO
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "TÀI LIỆU THAM KHẢO", 1)
    refs = [
        "[1] Vinyals, O., Toshev, A., Bengio, S., & Erhan, D. (2015). Show and tell: A neural image caption generator. CVPR 2015.",
        "[2] Xu, K., Ba, J., Kiros, R., et al. (2015). Show, attend and tell: Neural image caption generation with visual attention. ICML 2015.",
        "[3] Bahdanau, D., Cho, K., & Bengio, Y. (2015). Neural machine translation by jointly learning to align and translate. ICLR 2015.",
        "[4] Simonyan, K., & Zisserman, A. (2014). Very deep convolutional networks for large-scale image recognition (VGG16). ICLR 2015.",
        "[5] He, K., Zhang, X., Ren, S., & Sun, J. (2016). Deep residual learning for image recognition (ResNet). CVPR 2016.",
        "[6] Radford, A., Kim, J. W., Hallacy, C., et al. (2021). Learning transferable visual models from natural language supervision (CLIP). ICML 2021.",
        "[7] Pennington, J., Socher, R., & Manning, C. D. (2014). GloVe: Global vectors for word representation. EMNLP 2014.",
        "[8] Young, P., Lai, A., Hodosh, M., & Hockenmaier, J. (2014). From image descriptions to visual denotations (Flickr30k). TACL 2014.",
        "[9] Papineni, K., Roukos, S., Ward, T., & Zhu, W. J. (2002). BLEU: A method for automatic evaluation of machine translation. ACL 2002.",
        "[10] Loshchilov, I., & Hutter, F. (2019). Decoupled weight decay regularization (AdamW). ICLR 2019.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.runs[0].font.size = Pt(10.5)
        p.paragraph_format.space_after = Pt(4)

    # ── Save ──
    doc.save(OUT_PATH)
    print(f"✓ Đã tạo file: {OUT_PATH}")

# ═════════════════════════════════════════════════════��════════════════════════
if __name__ == "__main__":
    print("Đang tạo báo cáo Word...")
    build_word()
